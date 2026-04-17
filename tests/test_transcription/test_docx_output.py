# tests/test_transcription/test_docx_output.py
import os
import tempfile
from docx import Document
from autoso.transcription.docx_output import create_docx


def test_create_docx_returns_existing_file():
    path = create_docx("NS Exercise", "This is the full transcript text.")
    assert os.path.exists(path)
    assert path.endswith(".docx")
    os.unlink(path)


def test_create_docx_contains_title():
    path = create_docx("NS Exercise Transcript", "Some transcript content.")
    doc = Document(path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("NS Exercise Transcript" in h for h in headings)
    os.unlink(path)


def test_create_docx_contains_transcript_text():
    transcript = "SAF soldiers completed the exercise. NS training is important."
    path = create_docx("Title", transcript)
    doc = Document(path)
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "SAF soldiers completed the exercise" in full_text
    os.unlink(path)


def test_create_docx_writes_to_custom_path(tmp_path):
    out = str(tmp_path / "output.docx")
    result = create_docx("Title", "Text", output_path=out)
    assert result == out
    assert os.path.exists(out)
