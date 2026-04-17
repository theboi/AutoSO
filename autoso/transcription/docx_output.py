# autoso/transcription/docx_output.py
import tempfile
from typing import Optional

from docx import Document


def create_docx(
    title: str, transcript: str, output_path: Optional[str] = None
) -> str:
    """Create a DOCX file containing the transcript. Returns the file path."""
    if output_path is None:
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        output_path = tmp.name
        tmp.close()

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(transcript)
    doc.save(output_path)
    return output_path
